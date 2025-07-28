import logging
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, BaseLoader, Template
import markdown
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from io import BytesIO

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generate network documentation from parsed Visio data."""
    
    def __init__(self, template_dir: Path):
        self.template_dir = template_dir
        self.env = Environment(loader=FileSystemLoader(template_dir))
        self.professional_mode = True  # Default to professional templates
        
        # Add custom filters for templates
        self.env.filters['date_add_months'] = self._date_add_months
        
    def generate_documentation(self, diagram_data: Dict[str, Any], output_format: str = "html", ai_analysis: Dict[str, Any] = None, supplemental_data: Dict[str, Any] = None, template_config: Dict[str, Any] = None, organization_config: Dict[str, Any] = None) -> bytes:
        """
        Generate documentation from parsed diagram data.
        
        Args:
            diagram_data: Parsed Visio diagram data
            output_format: Output format (html, pdf, docx, markdown)
            ai_analysis: AI-enhanced content
            supplemental_data: User-provided supplemental information
            template_config: Template configuration
            organization_config: Organization branding configuration
            
        Returns:
            Generated document as bytes
        """
        logger.info(f"Generating {output_format} documentation with customization")
        
        # Process and enhance the diagram data
        enhanced_data = self._process_diagram_data(diagram_data)
        
        # Apply organization branding and configuration
        if organization_config:
            enhanced_data = self._apply_organization_branding(enhanced_data, organization_config)
        
        # Apply template configuration
        if template_config:
            enhanced_data = self._apply_template_config(enhanced_data, template_config)
        
        # Merge supplemental data if provided
        if supplemental_data:
            enhanced_data = self._merge_supplemental_data(enhanced_data, supplemental_data)
        
        # Merge AI analysis if provided
        if ai_analysis:
            enhanced_data['ai_analysis'] = ai_analysis
        
        if output_format == "html":
            return self._generate_html(enhanced_data)
        elif output_format == "pdf":
            return self._generate_pdf(enhanced_data)
        elif output_format == "docx":
            return self._generate_docx(enhanced_data)
        elif output_format == "markdown":
            return self._generate_markdown(enhanced_data)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _process_diagram_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Process and enhance diagram data with calculated metrics and statistics."""
        enhanced_data = data.copy()
        
        # Extract shapes and connections
        shapes = data.get("shapes", [])
        connections = data.get("connections", [])
        
        # Add professional documentation metadata
        enhanced_data.update({
            "project_name": data.get("project_name", "Network Infrastructure Implementation"),
            "customer_name": data.get("customer_name", ""),
            "doc_version": data.get("doc_version", "1.0"),
            "doc_status": data.get("doc_status", "Final"),
            "project_id": data.get("project_id", ""),
        })
        
        # Create device lookup for easier processing
        device_lookup = {shape["id"]: shape for shape in shapes}
        
        # Calculate device statistics
        device_types = self._get_device_types(shapes)
        device_type_distribution = self._get_device_type_distribution(shapes)
        devices_by_type = self._group_devices_by_type(shapes)
        
        # Calculate connection statistics
        connection_stats = self._calculate_connection_stats(shapes, connections, device_lookup)
        connections_enhanced = self._enhance_connections(connections, device_lookup)
        connection_types = self._get_connection_types(connections)
        
        # Calculate network metrics
        network_metrics = self._calculate_network_metrics(shapes, connections)
        
        # Identify network characteristics
        most_connected_devices = self._get_most_connected_devices(connection_stats, device_lookup)
        isolated_devices = self._get_isolated_devices(connection_stats, device_lookup)
        
        # Add all calculated data to enhanced_data
        enhanced_data.update({
            "device_types": device_types,
            "device_type_distribution": device_type_distribution,
            "device_type_counts": [device_type_distribution[dt] for dt in device_types],
            "devices_by_type": devices_by_type,
            "connections_enhanced": connections_enhanced,
            "connection_types": connection_types,
            "most_common_device_type": self._get_most_common_device_type(device_type_distribution),
            "avg_connections_per_device": network_metrics["avg_connections_per_device"],
            "network_density": network_metrics["network_density"],
            "most_connected_device": most_connected_devices[0] if most_connected_devices else {"name": "N/A", "connections": 0},
            "most_connected_devices": most_connected_devices,
            "isolated_devices": isolated_devices,
            "top_connected_devices_names": [d["name"] for d in most_connected_devices[:5]],
            "top_connected_devices_counts": [d["connections"] for d in most_connected_devices[:5]],
            "network_segments": self._estimate_network_segments(shapes, connections),
            "network_segments_list": self._analyze_network_segments(shapes, connections, device_lookup),
            "network_type": self._identify_network_type(shapes, connections),
            "topology_pattern": self._identify_topology_pattern(connection_stats),
            "redundancy_level": self._assess_redundancy_level(shapes, connections),
            "high_density_areas": self._identify_high_density_areas(connection_stats, device_lookup),
        })
        
        # Add connection counts to each device
        for shape in shapes:
            shape["connections_count"] = connection_stats.get(shape["id"], 0)
        
        return enhanced_data
    
    def _get_device_types(self, shapes: List[Dict]) -> List[str]:
        """Extract unique device types."""
        types = set()
        for shape in shapes:
            device_type = shape.get("type", "Unknown")
            types.add(device_type)
        return sorted(list(types))
    
    def _get_device_type_distribution(self, shapes: List[Dict]) -> Dict[str, int]:
        """Calculate distribution of devices by type."""
        distribution = defaultdict(int)
        for shape in shapes:
            device_type = shape.get("type", "Unknown")
            distribution[device_type] += 1
        return dict(distribution)
    
    def _group_devices_by_type(self, shapes: List[Dict]) -> Dict[str, List[Dict]]:
        """Group devices by their type."""
        grouped = defaultdict(list)
        for shape in shapes:
            device_type = shape.get("type", "Unknown")
            grouped[device_type].append(shape)
        return dict(grouped)
    
    def _calculate_connection_stats(self, shapes: List[Dict], connections: List[Dict], 
                                  device_lookup: Dict[str, Dict]) -> Dict[str, int]:
        """Calculate connection count for each device."""
        connection_count = defaultdict(int)
        
        for conn in connections:
            source_id = conn.get("source_id")
            target_id = conn.get("target_id")
            
            if source_id in device_lookup:
                connection_count[source_id] += 1
            if target_id in device_lookup:
                connection_count[target_id] += 1
                
        return dict(connection_count)
    
    def _enhance_connections(self, connections: List[Dict], device_lookup: Dict[str, Dict]) -> List[Dict]:
        """Enhance connections with device names."""
        enhanced = []
        
        for conn in connections:
            enhanced_conn = conn.copy()
            source_id = conn.get("source_id")
            target_id = conn.get("target_id")
            
            # Add device names
            source_device = device_lookup.get(source_id, {})
            target_device = device_lookup.get(target_id, {})
            
            enhanced_conn["source_name"] = source_device.get("name", source_id)
            enhanced_conn["target_name"] = target_device.get("name", target_id)
            
            enhanced.append(enhanced_conn)
            
        return enhanced
    
    def _get_connection_types(self, connections: List[Dict]) -> Dict[str, int]:
        """Count connections by type."""
        types = Counter()
        for conn in connections:
            conn_type = conn.get("connection_type", "Unknown")
            types[conn_type] += 1
        return dict(types)
    
    def _calculate_network_metrics(self, shapes: List[Dict], connections: List[Dict]) -> Dict[str, float]:
        """Calculate various network metrics."""
        num_devices = len(shapes)
        num_connections = len(connections)
        
        # Average connections per device
        avg_connections = (2 * num_connections / num_devices) if num_devices > 0 else 0
        
        # Network density (ratio of actual connections to possible connections)
        possible_connections = num_devices * (num_devices - 1) / 2
        density = (num_connections / possible_connections) if possible_connections > 0 else 0
        
        return {
            "avg_connections_per_device": avg_connections,
            "network_density": density,
            "total_devices": num_devices,
            "total_connections": num_connections
        }
    
    def _get_most_connected_devices(self, connection_stats: Dict[str, int], 
                                  device_lookup: Dict[str, Dict]) -> List[Dict]:
        """Get devices with the most connections."""
        devices_with_connections = []
        
        for device_id, conn_count in connection_stats.items():
            device = device_lookup.get(device_id, {})
            devices_with_connections.append({
                "id": device_id,
                "name": device.get("name", device_id),
                "type": device.get("type", "Unknown"),
                "connections": conn_count
            })
        
        # Sort by connection count (descending)
        devices_with_connections.sort(key=lambda x: x["connections"], reverse=True)
        
        return devices_with_connections
    
    def _get_isolated_devices(self, connection_stats: Dict[str, int], 
                            device_lookup: Dict[str, Dict]) -> List[Dict]:
        """Get devices with no connections."""
        isolated = []
        
        for device_id, device in device_lookup.items():
            if device_id not in connection_stats or connection_stats[device_id] == 0:
                isolated.append({
                    "id": device_id,
                    "name": device.get("name", device_id),
                    "type": device.get("type", "Unknown")
                })
                
        return isolated
    
    def _get_most_common_device_type(self, distribution: Dict[str, int]) -> str:
        """Get the most common device type."""
        if not distribution:
            return "N/A"
        return max(distribution.items(), key=lambda x: x[1])[0]
    
    def _estimate_network_segments(self, shapes: List[Dict], connections: List[Dict]) -> int:
        """Estimate the number of network segments (simplified)."""
        # This is a simplified estimation based on connection patterns
        # A more sophisticated approach would use graph algorithms
        if not shapes:
            return 0
        if not connections:
            return len(shapes)  # All devices are isolated
        
        # Simple heuristic: assume major segments based on connection density
        avg_connections = (2 * len(connections)) / len(shapes) if shapes else 0
        if avg_connections > 3:
            return 1  # Highly connected, likely single segment
        elif avg_connections > 1.5:
            return 2  # Moderately connected, possibly 2 segments
        else:
            return 3  # Sparsely connected, multiple segments
    
    def _analyze_network_segments(self, shapes: List[Dict], connections: List[Dict], 
                                device_lookup: Dict[str, Dict]) -> List[Dict]:
        """Analyze network segments in detail."""
        # Simplified segment analysis - in production, use graph algorithms
        segments = []
        
        # For now, return a simplified single segment
        if shapes and connections:
            segments.append({
                "name": "Main Network Segment",
                "device_count": len(shapes),
                "internal_connections": len(connections),
                "external_connections": 0,
                "key_devices": self._get_most_connected_devices(
                    self._calculate_connection_stats(shapes, connections, device_lookup), 
                    device_lookup
                )[:5]
            })
        
        return segments
    
    def _identify_network_type(self, shapes: List[Dict], connections: List[Dict]) -> str:
        """Identify the network type based on topology."""
        if not shapes:
            return "Empty"
        
        device_count = len(shapes)
        connection_count = len(connections)
        
        if connection_count == 0:
            return "Disconnected"
        
        avg_connections = (2 * connection_count) / device_count
        
        if avg_connections > 4:
            return "Mesh"
        elif avg_connections > 2.5:
            return "Hybrid"
        elif avg_connections > 1.5:
            return "Star"
        else:
            return "Bus"
    
    def _identify_topology_pattern(self, connection_stats: Dict[str, int]) -> str:
        """Identify the topology pattern."""
        if not connection_stats:
            return "None"
        
        connections = list(connection_stats.values())
        max_conn = max(connections)
        avg_conn = sum(connections) / len(connections)
        
        if max_conn > avg_conn * 3:
            return "Hub and Spoke"
        elif all(c >= 2 for c in connections):
            return "Redundant"
        else:
            return "Hierarchical"
    
    def _assess_redundancy_level(self, shapes: List[Dict], connections: List[Dict]) -> str:
        """Assess the redundancy level of the network."""
        if not shapes or not connections:
            return "None"
        
        avg_connections = (2 * len(connections)) / len(shapes)
        
        if avg_connections >= 3:
            return "High"
        elif avg_connections >= 2:
            return "Medium"
        elif avg_connections >= 1.5:
            return "Low"
        else:
            return "None"
    
    def _identify_high_density_areas(self, connection_stats: Dict[str, int], 
                                   device_lookup: Dict[str, Dict]) -> List[Dict]:
        """Identify areas with high connection density."""
        high_density = []
        
        if connection_stats:
            avg_connections = sum(connection_stats.values()) / len(connection_stats)
            threshold = avg_connections * 2
            
            for device_id, conn_count in connection_stats.items():
                if conn_count > threshold:
                    device = device_lookup.get(device_id, {})
                    high_density.append({
                        "id": device_id,
                        "name": device.get("name", device_id),
                        "type": device.get("type", "Unknown"),
                        "connections": conn_count
                    })
        
        return high_density
    
    def _generate_html(self, data: Dict[str, Any]) -> bytes:
        """Generate HTML documentation with custom template support."""
        # Check if custom template is provided
        if data.get("template", {}).get("html_template"):
            # Use custom template from database
            template_data = data["template"]
            template_str = template_data.get("html_template", "")
            css_styles = template_data.get("css_styles", "")
            header_template = template_data.get("header_template", "")
            footer_template = template_data.get("footer_template", "")
            
            # Create a Jinja2 template from string
            template = Template(template_str)
            
            # Prepare render data
            render_data = data.copy()
            render_data['css_styles'] = css_styles
            render_data['header_template'] = header_template
            render_data['footer_template'] = footer_template
            render_data['generated_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Add organization data if available
            if 'organization' in render_data:
                org = render_data['organization']
                # Create CSS variables for organization branding
                branding_css = f"""
                :root {{
                    --primary-color: {org.get('primary_color', '#1e3c72')};
                    --secondary-color: {org.get('secondary_color', '#2a5298')};
                    --accent-color: {org.get('accent_color', '#4CAF50')};
                    --font-family: {org.get('default_font_family', 'Arial, sans-serif')};
                    --font-size: {org.get('default_font_size', '14px')};
                }}
                """
                render_data['css_styles'] = branding_css + "\n" + css_styles
            
            # Ensure all required fields have defaults
            render_data.setdefault('title', 'Network Documentation')
            render_data.setdefault('project_name', data.get('project_name', 'Network Infrastructure'))
            render_data.setdefault('customer_name', data.get('customer_name', ''))
            render_data.setdefault('customer_organization', data.get('customer_organization', ''))
            render_data.setdefault('version', data.get('version', '1.0'))
            
            # Render template with data
            html_content = template.render(**render_data)
        else:
            # Use default template
            template_name = "professional_network_doc.html" if self.professional_mode else "network_doc.html"
            template = self.env.get_template(template_name)
            
            # Remove 'title' from data if it exists to avoid duplicate keyword argument
            render_data = data.copy()
            title = render_data.pop('title', 'Network Documentation')
            
            html_content = template.render(
                title=title,
                generated_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                diagram_data=data,
                **render_data  # Pass all enhanced data except title
            )
        
        return html_content.encode("utf-8")
    
    def _generate_pdf(self, data: Dict[str, Any]) -> bytes:
        """Generate PDF documentation."""
        # First generate HTML, then convert to PDF
        html_content = self._generate_html(data).decode("utf-8")
        
        try:
            # Import at function level to avoid conflicts
            import weasyprint
            import re
            
            # Remove script tags that might cause issues with WeasyPrint
            # Only remove script tags, keep all styling and content
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
            
            # Generate PDF with the full professional HTML template
            html_document = weasyprint.HTML(string=html_content)
            pdf_bytes = html_document.write_pdf()
            return pdf_bytes
        except Exception as e:
            logger.error(f"Error generating PDF with WeasyPrint: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise
    
    def _generate_docx(self, data: Dict[str, Any]) -> bytes:
        """Generate Word document with enhanced content."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set up document styles
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0x1e, 0x3c, 0x72)
        
        # Title page
        title_para = doc.add_paragraph()
        title_para.style = 'CustomTitle'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(data.get("title", "Network Documentation"))
        
        doc.add_paragraph()
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Subtitle')
        doc.add_paragraph(f"Source: {data.get('filename', 'Visio Diagram')}", style='Subtitle')
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            f"This network documentation provides a comprehensive overview of the network infrastructure. "
            f"The network consists of {len(data.get('shapes', []))} devices connected through "
            f"{len(data.get('connections', []))} connections."
        )
        
        # Key Statistics Table
        doc.add_heading("Key Statistics", level=2)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light List Accent 1'
        
        stats = [
            ("Total Devices", str(len(data.get('shapes', [])))),
            ("Total Connections", str(len(data.get('connections', [])))),
            ("Device Types", str(len(data.get('device_types', [])))),
            ("Most Common Device Type", data.get('most_common_device_type', 'N/A')),
            ("Average Connections per Device", f"{data.get('avg_connections_per_device', 0):.2f}"),
            ("Network Density", f"{data.get('network_density', 0):.3f}"),
            ("Isolated Devices", str(len(data.get('isolated_devices', []))))
        ]
        
        for i, (metric, value) in enumerate(stats):
            table.cell(i, 0).text = metric
            table.cell(i, 1).text = value
        
        doc.add_page_break()
        
        # Device Inventory
        doc.add_heading("Device Inventory", level=1)
        
        for device_type, devices in data.get('devices_by_type', {}).items():
            doc.add_heading(f"{device_type} ({len(devices)} devices)", level=2)
            
            for device in devices:
                doc.add_heading(device.get('name', 'Unknown'), level=3)
                doc.add_paragraph(f"ID: {device.get('id', 'N/A')}", style='List Bullet')
                doc.add_paragraph(f"Type: {device.get('type', 'Unknown')}", style='List Bullet')
                
                if device.get('connections_count'):
                    doc.add_paragraph(f"Connections: {device.get('connections_count')}", style='List Bullet')
                
                if device.get('properties'):
                    doc.add_paragraph("Properties:", style='List Bullet')
                    for key, value in device.get('properties', {}).items():
                        doc.add_paragraph(f"  • {key}: {value}", style='List Bullet 2')
        
        doc.add_page_break()
        
        # Network Connections
        doc.add_heading("Network Connections", level=1)
        
        # Connection summary
        doc.add_heading("Connection Types Summary", level=2)
        for conn_type, count in data.get('connection_types', {}).items():
            doc.add_paragraph(f"• {conn_type}: {count} connections", style='List Bullet')
        
        # Connection table
        doc.add_heading("Connection Details", level=2)
        if data.get('connections_enhanced'):
            conn_table = doc.add_table(rows=1, cols=4)
            conn_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = conn_table.rows[0].cells
            header_cells[0].text = 'Source Device'
            header_cells[1].text = 'Target Device'
            header_cells[2].text = 'Connection Type'
            header_cells[3].text = 'Properties'
            
            # Data rows
            for conn in data.get('connections_enhanced', [])[:50]:  # Limit to 50 for Word doc
                row_cells = conn_table.add_row().cells
                row_cells[0].text = conn.get('source_name', 'Unknown')
                row_cells[1].text = conn.get('target_name', 'Unknown')
                row_cells[2].text = conn.get('connection_type', 'Unknown')
                
                props = conn.get('properties', {})
                if props:
                    prop_text = ', '.join([f"{k}: {v}" for k, v in props.items()])
                    row_cells[3].text = prop_text
                else:
                    row_cells[3].text = '-'
        
        doc.add_page_break()
        
        # Network Analysis
        doc.add_heading("Network Analysis", level=1)
        
        # Most connected devices
        doc.add_heading("Highly Connected Devices", level=2)
        for device in data.get('most_connected_devices', [])[:10]:
            doc.add_paragraph(
                f"• {device['name']} ({device['type']}) - {device['connections']} connections",
                style='List Bullet'
            )
        
        # Isolated devices
        if data.get('isolated_devices'):
            doc.add_heading("Isolated Devices", level=2)
            for device in data.get('isolated_devices', []):
                doc.add_paragraph(f"• {device['name']} ({device['type']})", style='List Bullet')
        
        # Recommendations
        doc.add_heading("Recommendations", level=1)
        doc.add_paragraph(
            "Based on the network analysis, here are some observations and recommendations:"
        )
        
        recommendations = [
            "Regular network documentation updates are recommended to maintain accuracy.",
            "Consider implementing redundancy for highly connected devices to avoid single points of failure.",
            "Review isolated devices to determine if they should be connected or removed.",
            "Monitor network growth and plan for scalability as needed."
        ]
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def _generate_markdown(self, data: Dict[str, Any]) -> bytes:
        """Generate Markdown documentation."""
        # Choose template based on mode
        template_name = "professional_network_doc.md" if self.professional_mode else "network_doc.md"
        template = self.env.get_template(template_name)
        
        # Remove 'title' from data if it exists to avoid duplicate keyword argument
        render_data = data.copy()
        title = render_data.pop('title', 'Network Documentation')
        
        md_content = template.render(
            title=title,
            generated_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            diagram_data=data,
            **render_data  # Pass all enhanced data except title
        )
        return md_content.encode("utf-8")
    
    def _merge_supplemental_data(self, data: Dict[str, Any], supplemental: Dict[str, Any]) -> Dict[str, Any]:
        """Merge supplemental information into the main data"""
        merged = data.copy()
        answers = supplemental.get("answers", {})
        
        # Apply network design pattern (handle "not sure" with AI inference)
        if "network_design" in answers:
            design_choice = answers["network_design"]
            if design_choice == "not_sure":
                # Mark for AI analysis - will be handled by LLM suggestions
                merged["network_design_pattern"] = "AI-Inferred"
                merged["network_type"] = "AI-Analyzed Topology"
                merged["requires_ai_analysis"] = True
            else:
                merged["network_design_pattern"] = design_choice
                merged["network_type"] = self._map_design_to_type(design_choice)
        
        # Apply VLAN information
        if "vlan_list" in answers and answers["vlan_list"]:
            merged["vlans"] = self._parse_vlan_list(answers["vlan_list"])
        
        # Apply device details (handle "not sure" entries)
        if "device_details" in answers and answers["device_details"]:
            merged = self._merge_device_details(merged, answers["device_details"])
        
        # Apply port channels
        if "port_channels" in answers and answers["port_channels"]:
            merged["port_channels"] = answers["port_channels"]
        
        # Apply site details
        if "site_details" in answers:
            merged["site_information"] = answers["site_details"]
        
        # Apply logical connections
        if "missing_connections" in answers and answers["missing_connections"]:
            merged["logical_connections"] = answers["missing_connections"]
        
        # Track user-provided vs AI-inferred information
        merged["supplemental_info_source"] = {
            "user_provided": len([k for k, v in answers.items() if v and v != "not_sure"]),
            "ai_inferred": len([k for k, v in answers.items() if v == "not_sure"]),
            "total_questions": len(answers)
        }
        
        return merged
    
    def _map_design_to_type(self, design: str) -> str:
        """Map design pattern to network type"""
        mapping = {
            "collapsed_core": "Collapsed Core Architecture",
            "three_tier": "Three-Tier Hierarchical",
            "spine_leaf": "Spine-Leaf Data Center",
            "hub_spoke": "Hub and Spoke",
            "mesh": "Mesh Topology",
            "star": "Star Topology",
            "ring": "Ring Topology", 
            "bus": "Bus Topology",
            "hybrid": "Hybrid Architecture",
            "dmz": "DMZ/Security-Focused",
            "campus": "Campus Network",
            "wan": "WAN/Branch Network",
            "not_sure": "AI-Analyzed Topology"
        }
        return mapping.get(design, "Custom Architecture")
    
    def _parse_vlan_list(self, vlan_text: str) -> List[Dict[str, str]]:
        """Parse VLAN information from text"""
        vlans = []
        lines = vlan_text.strip().split('\n')
        
        for line in lines:
            parts = [p.strip() for p in line.split(',')]
            if len(parts) >= 2:
                vlans.append({
                    "id": parts[0],
                    "name": parts[1],
                    "description": parts[2] if len(parts) > 2 else ""
                })
        
        return vlans
    
    def _merge_device_details(self, data: Dict[str, Any], device_details: List[Dict]) -> Dict[str, Any]:
        """Merge additional device details, handling 'not sure' responses"""
        # Create lookup by device name
        detail_lookup = {d.get("name"): d for d in device_details if d.get("name")}
        
        # Track devices needing AI inference
        devices_needing_inference = []
        
        # Update shapes with additional details
        for shape in data.get("shapes", []):
            name = shape.get("name")
            if name and name in detail_lookup:
                details = detail_lookup[name]
                device_needs_inference = False
                
                # Handle model information
                if "model" in details:
                    if details["model"] == "Not Sure":
                        shape["model_needs_inference"] = True
                        device_needs_inference = True
                    else:
                        shape["model"] = details["model"]
                
                # Handle IP information
                if "ip" in details and details["ip"]:
                    shape.setdefault("properties", {})["ip_address"] = details["ip"]
                
                # Handle role information
                if "role" in details:
                    if details["role"] == "Not Sure":
                        shape["role_needs_inference"] = True
                        device_needs_inference = True
                    else:
                        shape["role"] = details["role"]
                
                if device_needs_inference:
                    devices_needing_inference.append(name)
        
        # Track devices that need AI inference
        if devices_needing_inference:
            data["devices_needing_ai_inference"] = devices_needing_inference
        
        return data
    
    def _apply_organization_branding(self, data: Dict[str, Any], organization_config: Dict[str, Any]) -> Dict[str, Any]:
        """Apply organization branding and configuration to document data"""
        enhanced_data = data.copy()
        
        # Organization Information
        enhanced_data.update({
            "organization_name": organization_config.get("name", ""),
            "organization_display_name": organization_config.get("display_name", organization_config.get("name", "")),
            "organization_contact": organization_config.get("primary_contact", ""),
            "organization_email": organization_config.get("email", ""),
            "organization_phone": organization_config.get("phone", ""),
            "organization_website": organization_config.get("website", ""),
            "organization_address": self._format_organization_address(organization_config),
            "logo_url": organization_config.get("logo_url", ""),
        })
        
        # Branding Configuration
        branding = {
            "primary_color": organization_config.get("primary_color", "#1e3c72"),
            "secondary_color": organization_config.get("secondary_color", "#2a5298"),
            "accent_color": organization_config.get("accent_color", "#4CAF50"),
            "font_family": organization_config.get("default_font_family", "Arial"),
            "font_size": organization_config.get("default_font_size", "14px"),
            "letterhead_html": organization_config.get("letterhead_html", ""),
            "footer_html": organization_config.get("footer_html", ""),
            "document_numbering_format": organization_config.get("document_numbering_format", "DOC-{year}-{seq:04d}")
        }
        
        enhanced_data["branding"] = branding
        enhanced_data["custom_branding"] = organization_config.get("branding_config", {})
        
        return enhanced_data
    
    def _apply_template_config(self, data: Dict[str, Any], template_config: Dict[str, Any]) -> Dict[str, Any]:
        """Apply template-specific configuration to document data"""
        enhanced_data = data.copy()
        
        # Template Configuration
        template_settings = {
            "template_name": template_config.get("name", "Default Template"),
            "template_type": template_config.get("template_type", "network_documentation"),
            "template_version": template_config.get("version", "1.0"),
            "supported_formats": template_config.get("supported_formats", ["html", "pdf"]),
            "page_margins": template_config.get("page_margins", {"top": "1in", "right": "1in", "bottom": "1in", "left": "1in"}),
            "font_config": template_config.get("font_config", {}),
            "color_scheme": template_config.get("color_scheme", {}),
            "logo_config": template_config.get("logo_config", {}),
            "section_config": template_config.get("section_config", {})
        }
        
        enhanced_data["template"] = template_settings
        
        # Override branding with template-specific settings if provided
        if template_config.get("color_scheme"):
            enhanced_data.setdefault("branding", {}).update(template_config["color_scheme"])
        
        if template_config.get("font_config"):
            font_config = template_config["font_config"]
            branding = enhanced_data.setdefault("branding", {})
            branding["font_family"] = font_config.get("primary_font", branding.get("font_family", "Arial"))
            branding["font_size"] = font_config.get("base_size", branding.get("font_size", "14px"))
        
        return enhanced_data
    
    def _format_organization_address(self, organization_config: Dict[str, Any]) -> str:
        """Format organization address for display"""
        address_parts = []
        
        if organization_config.get("address_line1"):
            address_parts.append(organization_config["address_line1"])
        
        if organization_config.get("address_line2"):
            address_parts.append(organization_config["address_line2"])
        
        city_state_zip = []
        if organization_config.get("city"):
            city_state_zip.append(organization_config["city"])
        
        if organization_config.get("state"):
            city_state_zip.append(organization_config["state"])
        
        if organization_config.get("postal_code"):
            city_state_zip.append(organization_config["postal_code"])
        
        if city_state_zip:
            address_parts.append(", ".join(city_state_zip))
        
        if organization_config.get("country"):
            address_parts.append(organization_config["country"])
        
        return "\n".join(address_parts)
    
    def _date_add_months(self, date_str: str, months: int) -> str:
        """Add months to a date string for template use."""
        try:
            # Parse the date string
            date = datetime.strptime(date_str.split()[0], "%Y-%m-%d")
            # Simple month addition (approximate)
            new_date = date + timedelta(days=30 * months)
            return new_date.strftime("%Y-%m-%d")
        except:
            # If parsing fails, return original
            return date_str